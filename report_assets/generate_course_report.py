from __future__ import annotations

import sqlite3
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "report_assets"
IMG_DIR = ASSETS / "requirements.files"
PROCESSED = ASSETS / "processed"
OUT_FILE = ROOT / "移动终端程序设计结课报告-张善茹-实时聊天系统.docx"

CN_FONT = "宋体"
HEI_FONT = "黑体"
EN_FONT = "Times New Roman"
CAPTION_SIZE = 10.5
BODY_SIZE = 12


def font_path(*names: str) -> str | None:
    font_dir = Path("C:/Windows/Fonts")
    for name in names:
        p = font_dir / name
        if p.exists():
            return str(p)
    return None


FONT_HEI = font_path("simhei.ttf", "msyh.ttc", "simsun.ttc")
FONT_SONG = font_path("simsun.ttc", "msyh.ttc", "simhei.ttf")
FONT_EN = font_path("times.ttf", "arial.ttf")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = FONT_HEI if bold else FONT_SONG
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def ensure_dirs() -> None:
    PROCESSED.mkdir(parents=True, exist_ok=True)


def prepare_image(name: str, max_width: int = 1200, jpeg_quality: int = 92) -> Path:
    src = IMG_DIR / name
    out = PROCESSED / (src.stem + ".jpg")
    if out.exists() and out.stat().st_mtime >= src.stat().st_mtime:
        return out
    img = Image.open(src).convert("RGB")
    if img.width > max_width:
        ratio = max_width / img.width
        img = img.resize((max_width, int(img.height * ratio)), Image.Resampling.LANCZOS)
    img.save(out, quality=jpeg_quality, optimize=True)
    return out


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    """Wrap mixed Chinese/English text by measured pixel width."""
    lines: list[str] = []
    for raw in text.split("\n"):
        raw = raw.strip()
        if not raw:
            lines.append("")
            continue
        if text_size(draw, raw, font)[0] <= max_width:
            lines.append(raw)
            continue

        words = raw.split(" ") if " " in raw else list(raw)
        line = ""
        for word in words:
            piece = word if not line or " " not in raw else " " + word
            trial = line + piece
            if text_size(draw, trial, font)[0] <= max_width:
                line = trial
                continue
            if line:
                lines.append(line.strip())
                line = word
            else:
                # Extremely long token: split by character so it still stays inside the box.
                token = ""
                for ch in word:
                    trial_token = token + ch
                    if text_size(draw, trial_token, font)[0] <= max_width:
                        token = trial_token
                    else:
                        if token:
                            lines.append(token)
                        token = ch
                line = token
        if line:
            lines.append(line.strip())
    return lines


def draw_round_rect(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str, width: int = 3, radius: int = 22) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_card(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: str,
    outline: str,
    title_color: str = "#111827",
    body_color: str = "#374151",
    title_size: int = 28,
    body_size: int = 22,
    align: str = "center",
) -> None:
    draw_round_rect(draw, box, fill, outline, width=3, radius=22)
    x1, y1, x2, y2 = box
    pad_x, pad_y = 28, 24
    max_width = x2 - x1 - pad_x * 2
    max_height = y2 - y1 - pad_y * 2

    for shrink in range(0, 8):
        title_font = pil_font(title_size - shrink, bold=True)
        body_font = pil_font(body_size - shrink)
        title_lines = wrap_text(draw, title, title_font, max_width)
        body_lines = wrap_text(draw, body, body_font, max_width)
        title_h = sum(text_size(draw, line, title_font)[1] for line in title_lines) + max(0, len(title_lines) - 1) * 7
        body_h = sum(text_size(draw, line, body_font)[1] for line in body_lines) + max(0, len(body_lines) - 1) * 7
        total_h = title_h + (14 if body else 0) + body_h
        if total_h <= max_height:
            break

    y = y1 + (y2 - y1 - total_h) / 2
    for line in title_lines:
        w, h = text_size(draw, line, title_font)
        x = x1 + pad_x if align == "left" else x1 + (x2 - x1 - w) / 2
        draw.text((x, y), line, font=title_font, fill=title_color)
        y += h + 7
    if body:
        y += 7
    for line in body_lines:
        w, h = text_size(draw, line, body_font)
        x = x1 + pad_x if align == "left" else x1 + (x2 - x1 - w) / 2
        draw.text((x, y), line, font=body_font, fill=body_color)
        y += h + 7


def draw_label(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    fill: str,
    size: int = 20,
    bg: str = "#ffffff",
) -> None:
    font = pil_font(size)
    x, y = xy
    w, h = text_size(draw, text, font)
    draw.rounded_rectangle((x - 10, y - 6, x + w + 10, y + h + 8), radius=10, fill=bg)
    draw.text((x, y), text, font=font, fill=fill)


def arrow_head(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str, size: int = 15) -> None:
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - direction * size, ey - size * 0.6), (ex - direction * size, ey + size * 0.6)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - size * 0.6, ey - direction * size), (ex + size * 0.6, ey - direction * size)]
    draw.polygon(pts, fill=fill)


def polyline_arrow(draw: ImageDraw.ImageDraw, points: Sequence[tuple[int, int]], fill: str = "#4b5563", width: int = 4) -> None:
    for start, end in zip(points, points[1:]):
        draw.line([start, end], fill=fill, width=width)
    arrow_head(draw, points[-2], points[-1], fill)


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str, width: int) -> None:
    title_font = pil_font(44, bold=True)
    sub_font = pil_font(22)
    draw.text((70, 50), title, font=title_font, fill="#111827")
    y = 112
    for line in wrap_text(draw, subtitle, sub_font, width - 140):
        draw.text((70, y), line, font=sub_font, fill="#4b5563")
        y += 32


def draw_lane(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, fill: str, outline: str) -> None:
    draw.rounded_rectangle(box, radius=28, fill=fill, outline=outline, width=2)
    x1, y1, _, _ = box
    font = pil_font(24, bold=True)
    draw.text((x1 + 24, y1 + 22), title, font=font, fill=outline)


def make_architecture_diagram() -> Path:
    out = PROCESSED / "system_architecture.png"
    img = Image.new("RGB", (1800, 1050), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "系统总体架构",
        "前端负责移动端页面与交互，后端同时提供 REST API 和 Socket.io 实时通信，数据库层通过 Sequelize 将业务数据持久化到 SQLite。",
        img.width,
    )

    draw_lane(draw, (70, 185, 555, 960), "前端层", "#fffaf0", "#d97706")
    draw_lane(draw, (650, 185, 1245, 960), "后端服务层", "#f8fafc", "#475569")
    draw_lane(draw, (1345, 185, 1730, 960), "数据层", "#fff5f5", "#dc2626")

    draw_card(draw, (125, 275, 510, 425), "移动端页面", "Auth / Home / ChatBox\nRoomList / ContactList", "#ffffff", "#f59e0b")
    draw_card(draw, (125, 505, 510, 655), "前端状态", "localStorage 保存 token\nRoomContext 记录当前会话", "#ffffff", "#f59e0b")
    draw_card(draw, (125, 735, 510, 885), "Socket.io Client", "join_room\nsend_message\nreceive_message", "#ffffff", "#16a34a")

    draw_card(draw, (710, 250, 1190, 395), "Express REST API", "/api/auth /api/room\n/api/contact /api/private", "#ffffff", "#2563eb")
    draw_card(draw, (710, 485, 1190, 630), "业务控制器", "认证、房间、联系人\n历史消息查询", "#ffffff", "#475569")
    draw_card(draw, (710, 720, 1190, 865), "Socket.io Server", "鉴权中间件\n加入房间与消息广播", "#ffffff", "#16a34a")

    draw_card(draw, (1390, 335, 1685, 500), "Sequelize Models", "User / Room\nRoomMessage / PrivateMessage\nContact", "#ffffff", "#b91c1c", title_size=25, body_size=20)
    draw_card(draw, (1390, 650, 1685, 845), "SQLite 数据库", "users / rooms / RoomUser\ncontacts / room_messages\nprivate_messages", "#ffffff", "#dc2626", title_size=25, body_size=20)

    polyline_arrow(draw, [(510, 350), (710, 322)], "#2563eb", 5)
    draw_label(draw, (565, 286), "HTTP / JSON", "#2563eb")
    polyline_arrow(draw, [(510, 580), (625, 580), (625, 558), (710, 558)], "#475569", 5)
    draw_label(draw, (548, 610), "token 与当前用户", "#475569")
    polyline_arrow(draw, [(510, 810), (710, 792)], "#16a34a", 5)
    draw_label(draw, (552, 825), "实时事件", "#16a34a")

    polyline_arrow(draw, [(1190, 322), (1290, 322), (1290, 405), (1390, 405)], "#2563eb", 5)
    draw_label(draw, (1245, 285), "查询 / 写入", "#2563eb")
    polyline_arrow(draw, [(1190, 558), (1290, 558), (1290, 438), (1390, 438)], "#475569", 5)
    draw_label(draw, (1210, 590), "ORM 调用", "#475569")
    polyline_arrow(draw, [(1190, 792), (1290, 792), (1290, 470), (1390, 470)], "#16a34a", 5)
    draw_label(draw, (1215, 815), "消息落库", "#16a34a")
    polyline_arrow(draw, [(1538, 500), (1538, 650)], "#dc2626", 5)
    draw_label(draw, (1560, 560), "同步数据表", "#dc2626")

    note_font = pil_font(19)
    note = "说明：REST 路径适合登录、列表、历史记录等请求；Socket 路径适合聊天消息、房间通知等实时事件。两条路径最终都通过模型层访问同一个 SQLite 数据库。"
    y = 980
    for line in wrap_text(draw, note, note_font, img.width - 160):
        draw.text((80, y), line, font=note_font, fill="#64748b")
        y += 28
    img.save(out)
    return out


def make_message_flow_diagram() -> Path:
    out = PROCESSED / "message_flow.png"
    img = Image.new("RGB", (1800, 1180), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "消息发送与持久化流程",
        "以用户在 ChatBox 中发送一条消息为例，流程同时包含前端状态更新、Socket.io 事件转发和 SQLite 数据库存储。",
        img.width,
    )

    lanes = [
        ((90, 185, 1710, 430), "前端：React ChatBox", "#fffaf0", "#d97706"),
        ((90, 470, 1710, 735), "后端：Socket.io + 业务逻辑", "#f8fafc", "#475569"),
        ((90, 775, 1710, 990), "数据层：Sequelize + SQLite", "#fff5f5", "#dc2626"),
    ]
    for box, title, fill, outline in lanes:
        draw_lane(draw, box, title, fill, outline)

    cards = {
        "input": (190, 255, 475, 385, "1 输入消息", "填写消息内容\n点击 Send 按钮", "#f59e0b"),
        "payload": (600, 255, 910, 385, "2 组装数据", "roomId、senderId、content\n判断 isPrivate", "#f59e0b"),
        "update": (1320, 255, 1615, 385, "7 更新页面", "追加 messageList\n自动滚动到底部", "#f59e0b"),
        "receive": (190, 545, 475, 675, "3 接收事件", "监听 send_message\n读取 messageData", "#475569"),
        "decide": (600, 545, 910, 675, "4 判断类型", "私聊写 PrivateMessage\n群聊写 RoomMessage", "#475569"),
        "broadcast": (1320, 545, 1615, 675, "6 房间广播", "io.to(roomId).emit\nreceive_message", "#16a34a"),
        "save": (600, 825, 910, 945, "5 持久化", "Sequelize create()\n写入 SQLite", "#dc2626"),
    }
    for box in cards.values():
        x1, y1, x2, y2, title, body, color = box
        fill = "#ffffff"
        draw_card(draw, (x1, y1, x2, y2), title, body, fill, color, title_size=24, body_size=19)

    polyline_arrow(draw, [(475, 320), (600, 320)], "#d97706", 5)
    polyline_arrow(draw, [(755, 385), (755, 455), (332, 455), (332, 545)], "#2563eb", 5)
    draw_label(draw, (800, 420), "emit('send_message')", "#2563eb")
    polyline_arrow(draw, [(475, 610), (600, 610)], "#475569", 5)
    polyline_arrow(draw, [(755, 675), (755, 825)], "#dc2626", 5)
    draw_label(draw, (780, 742), "写库", "#dc2626")
    polyline_arrow(draw, [(910, 885), (1220, 885), (1220, 610), (1320, 610)], "#16a34a", 5)
    draw_label(draw, (1080, 845), "保存后广播", "#16a34a")
    polyline_arrow(draw, [(1468, 545), (1468, 385)], "#16a34a", 5)
    draw_label(draw, (1490, 445), "receive_message", "#16a34a")

    note_font = pil_font(20)
    note = "设计要点：实时消息先通过 Socket 到达后端，再由后端统一写库和广播；页面切换或刷新时，前端再通过 REST 接口读取历史消息，因此“即时显示”和“可追溯记录”保持一致。"
    y = 1032
    for line in wrap_text(draw, note, note_font, img.width - 180):
        draw.text((95, y), line, font=note_font, fill="#475569")
        y += 32
    img.save(out)
    return out


def make_mobile_workflow_diagram() -> Path:
    out = PROCESSED / "mobile_workflow.png"
    img = Image.new("RGB", (1700, 760), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "移动端主要操作流程",
        "从用户第一次进入应用到开始聊天，界面围绕“认证、找房间、找联系人、发消息”四类动作组织。",
        img.width,
    )
    steps = [
        ("1 注册 / 登录", "填写邮箱和密码\n获取 JWT token", "#2563eb"),
        ("2 进入首页", "导航切换 Rooms\nContacts / Find people", "#d97706"),
        ("3 房间操作", "创建、加入\n查看我的房间", "#d97706"),
        ("4 公共聊天", "选择房间\n发送实时消息", "#16a34a"),
        ("5 联系人与私聊", "添加联系人\n进入一对一会话", "#7c3aed"),
    ]
    x, y = 120, 275
    for i, (title, body, color) in enumerate(steps):
        box = (x + i * 310, y, x + i * 310 + 245, y + 210)
        draw_card(draw, box, title, body, "#ffffff", color, title_size=24, body_size=19)
        if i < len(steps) - 1:
            polyline_arrow(draw, [(box[2], y + 105), (box[2] + 65, y + 105)], "#64748b", 4)
    draw.rounded_rectangle((130, 560, 1560, 650), radius=22, fill="#f8fafc", outline="#cbd5e1", width=2)
    font = pil_font(21)
    tips = "移动端设计思路：把高频动作放在顶部导航和卡片按钮中，尽量减少页面跳转；聊天输入区固定在聊天窗口底部，符合手机聊天应用的使用习惯。"
    y2 = 585
    for line in wrap_text(draw, tips, font, 1360):
        draw.text((165, y2), line, font=font, fill="#475569")
        y2 += 30
    img.save(out)
    return out


def make_database_diagram() -> Path:
    out = PROCESSED / "database_relationship.png"
    img = Image.new("RGB", (1700, 1020), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw_title(
        draw,
        "数据库关系设计",
        "数据库以用户为核心，围绕聊天室、联系人、公共消息和私聊消息建立关系，既支持多人聊天室，也支持一对一私聊。",
        img.width,
    )

    entities = {
        "users": ((100, 235, 465, 455), "users", "id 主键\nusername / email\npassword 哈希\nstatus 在线状态", "#2563eb"),
        "rooms": ((1190, 235, 1555, 455), "rooms", "id 主键\nid_admin 创建者\nroomName\n说明 description", "#d97706"),
        "roomuser": ((650, 250, 1035, 430), "RoomUser", "userId\nroomId\n用户与房间多对多关系", "#64748b"),
        "contacts": ((100, 650, 465, 850), "contacts", "id_user\nid_contact\n双向联系人关系", "#7c3aed"),
        "private": ((650, 650, 1035, 850), "private_messages", "senderId\nreceiverId\ncontent\ncreatedAt", "#16a34a"),
        "roommsg": ((1190, 650, 1555, 850), "room_messages", "senderId\nroomId\ncontent\ncreatedAt", "#dc2626"),
    }
    for box, title, body, color in entities.values():
        draw_card(draw, box, title, body, "#ffffff", color, title_size=26, body_size=20, align="left")

    polyline_arrow(draw, [(465, 345), (650, 340)], "#64748b", 4)
    draw_label(draw, (520, 300), "1 : N", "#64748b")
    polyline_arrow(draw, [(1035, 340), (1190, 345)], "#64748b", 4)
    draw_label(draw, (1080, 300), "N : 1", "#64748b")
    polyline_arrow(draw, [(282, 455), (282, 650)], "#7c3aed", 4)
    draw_label(draw, (305, 535), "联系人", "#7c3aed")
    polyline_arrow(draw, [(465, 760), (650, 760)], "#16a34a", 4)
    draw_label(draw, (520, 725), "私聊双方", "#16a34a")
    polyline_arrow(draw, [(1372, 455), (1372, 650)], "#dc2626", 4)
    draw_label(draw, (1395, 535), "房间消息", "#dc2626")
    polyline_arrow(draw, [(465, 410), (560, 410), (560, 705), (650, 705)], "#16a34a", 4)
    draw_label(draw, (500, 455), "发送者/接收者", "#16a34a")
    polyline_arrow(draw, [(465, 430), (530, 430), (530, 600), (1110, 600), (1110, 745), (1190, 745)], "#dc2626", 4)
    draw_label(draw, (980, 565), "公共消息发送者", "#dc2626")

    note_font = pil_font(19)
    note = "说明：RoomUser 解决用户与聊天室的多对多关系；room_messages 保存公共聊天室消息；private_messages 保存两名用户之间的私聊消息；contacts 用于构建联系人列表。"
    y = 930
    for line in wrap_text(draw, note, note_font, img.width - 170):
        draw.text((100, y), line, font=note_font, fill="#475569")
        y += 28
    img.save(out)
    return out


def set_run_font(run, size: float = BODY_SIZE, bold: bool = False, italic: bool = False, color: str | None = None, east_asia: str = CN_FONT, latin: str = EN_FONT) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, first_line: bool = False, line_spacing: float = 1.5, before: float = 0, after: float = 0) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)


def add_para(doc: Document, text: str = "", first_line: bool = True, align: int | None = None, size: float = BODY_SIZE, bold: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    before_after = 6
    set_paragraph_format(p, first_line=False, before=before_after, after=before_after)
    if level == 1:
        size, font, bold = 16, HEI_FONT, True
    elif level == 2:
        size, font, bold = 14, CN_FONT, True
    else:
        size, font, bold = 12, CN_FONT, False
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, east_asia=font)


def add_caption(doc: Document, text: str, kind: str = "图") -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, line_spacing=1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=CAPTION_SIZE, east_asia=CN_FONT, latin=EN_FONT)


def add_picture(doc: Document, path: Path, width_in: float, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_in))
    if caption:
        add_caption(doc, caption)


def add_picture_grid(doc: Document, images: Sequence[tuple[Path, str]], caption: str, width_in: float = 2.25) -> None:
    table = doc.add_table(rows=1, cols=len(images))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, (path, label) in enumerate(images):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(width_in))
        label_p = cell.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(label_p, first_line=False, line_spacing=1.0)
        r = label_p.add_run(label)
        set_run_font(r, size=9, east_asia=CN_FONT)
    remove_table_borders(table)
    add_caption(doc, caption)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:" + edge
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        if edge_data is None:
            element.set(qn("w:val"), "nil")
        else:
            for key, value in edge_data.items():
                element.set(qn("w:" + key), str(value))


def three_line_table(table) -> None:
    line = {"val": "single", "sz": "4", "space": "0", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=line, bottom=line)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=line)


def set_cell_text(cell, text: str, size: float = 10.5, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_format(p, first_line=False, line_spacing=1.0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_table_caption(doc: Document, text: str) -> None:
    add_caption(doc, text, kind="表")


def add_simple_table(doc: Document, caption: str, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None) -> None:
    add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        cells = table.add_row().cells
        for i, item in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.RIGHT if item.replace(".", "", 1).isdigit() else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], item, align=align)
    three_line_table(table)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)


def db_rows() -> list[tuple[str, str, str]]:
    db = ROOT / "server" / "src" / "database" / "database.sqlite"
    if not db.exists():
        return []
    con = sqlite3.connect(db)
    cur = con.cursor()
    tables = [r[0] for r in cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name")]
    rows: list[tuple[str, str, str]] = []
    aliases = {
        "users": "用户账号、邮箱、加密密码、在线状态",
        "rooms": "聊天室名称、描述、创建者",
        "RoomUser": "用户与聊天室的多对多关系",
        "room_messages": "公共聊天室消息记录",
        "private_messages": "私聊消息记录",
        "contacts": "联系人关系",
    }
    for table in tables:
        count = cur.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0]
        cols = cur.execute(f'PRAGMA table_info("{table}")').fetchall()
        fields = "、".join(c[1] for c in cols[:6])
        if len(cols) > 6:
            fields += "等"
        rows.append((table, aliases.get(table, "业务数据表"), f"{fields}；当前记录数：{count}"))
    con.close()
    return rows


def add_code_listing(doc: Document, caption: str, code: str) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_border(cell, top={"val": "single", "sz": "4", "color": "DDDDDD"}, bottom={"val": "single", "sz": "4", "color": "DDDDDD"}, left={"val": "single", "sz": "4", "color": "DDDDDD"}, right={"val": "single", "sz": "4", "color": "DDDDDD"})
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F7F7F7")
    cell._tc.get_or_add_tcPr().append(shading)
    cell.text = ""
    lines = code.strip("\n").splitlines()
    for index, line in enumerate(lines, start=1):
        p = cell.paragraphs[0] if index == 1 else cell.add_paragraph()
        set_paragraph_format(p, first_line=False, line_spacing=1.5)
        run = p.add_run(f"{index}. {line}")
        set_run_font(run, size=10.5, east_asia=CN_FONT, latin=EN_FONT)


def add_algorithm(doc: Document) -> None:
    add_caption(doc, "算法4.1 实时消息发送与同步显示流程")
    lines = [
        "输入：roomId、message、currentUser",
        "输出：数据库中的消息记录、所有在线会话中的最新消息",
        "1: if message 为空 then return",
        "2: 根据 roomId 判断当前会话是公共聊天室还是私聊",
        "3: 前端通过 socket.emit('send_message', messageData) 发送消息",
        "4: 后端接收事件后，若 isPrivate 为真则写入 PrivateMessage，否则写入 RoomMessage",
        "5: 后端执行 io.to(roomId).emit('receive_message', newMessage)",
        "6: 前端监听 receive_message，将新消息追加到 messageList 并自动滚动到底部",
        "7: return 最新聊天界面",
    ]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_border(cell, top={"val": "single", "sz": "4", "color": "000000"}, bottom={"val": "single", "sz": "4", "color": "000000"})
    cell.text = ""
    for line in lines:
        p = cell.add_paragraph()
        set_paragraph_format(p, first_line=False, line_spacing=1.5)
        r = p.add_run(line)
        set_run_font(r, size=10.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    style.font.size = Pt(BODY_SIZE)


def add_cover(doc: Document, logo: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo), width=Inches(1.25))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    r = title.add_run("移动终端程序设计结课报告")
    set_run_font(r, size=22, bold=True, east_asia=HEI_FONT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    r = subtitle.add_run("基于 React + Socket.io 的移动端实时聊天系统")
    set_run_font(r, size=16, bold=True, east_asia=CN_FONT)

    info = [
        ("班    级：", "________________"),
        ("学    号：", "________________"),
        ("姓    名：", "张善茹"),
        ("完成日期：", "2026 年 6 月 21 日"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    for i, (k, v) in enumerate(info):
        set_cell_text(table.cell(i, 0), k, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i, 1), v, size=14, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("东北大学秦皇岛分校计算机与通信工程学院")
    set_run_font(r, size=14, bold=True, east_asia=CN_FONT)
    doc.add_page_break()


def add_score_table(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("课程评分表")
    set_run_font(r, size=16, bold=True, east_asia=HEI_FONT)
    rows = [
        ("1", "界面要求：美观、友好，不能过于粗糙，具备一定设计美感（20分）", ""),
        ("2", "功能要求：基本功能实现，需要可以对功能进行扩充（20分）", ""),
        ("3", "对数据的持久化保存方式：可以使用文件，也可以使用数据库（20分）", ""),
        ("4", "代码要求：书写规范，结构合理，结合必要的注释（20分）", ""),
        ("5", "文档要求：内容详细，结构合理，格式正确（20分）", ""),
        ("合计", "", ""),
    ]
    add_simple_table(doc, "表0.1 结课报告评分表", ("序号", "要求", "得分"), rows, widths=(1.6, 12.0, 2.2))
    add_para(doc, "评阅人：____________", first_line=False)
    add_para(doc, "日期：____________", first_line=False)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_run_font(r, size=16, bold=True, east_asia=HEI_FONT)
    items = [
        "1 项目概述",
        "  1.1 选题背景与目标",
        "  1.2 开发环境与技术栈",
        "  1.3 系统总体架构",
        "2 需求分析与总体设计",
        "  2.1 功能需求",
        "  2.2 界面与交互设计",
        "  2.3 数据库设计",
        "3 系统功能实现",
        "  3.1 用户注册与登录",
        "  3.2 聊天室创建、加入与切换",
        "  3.3 公共聊天与私聊",
        "  3.4 前后端运行效果",
        "4 核心实现与代码说明",
        "5 测试结果与问题分析",
        "6 总结与心得",
    ]
    for item in items:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, line_spacing=1.5)
        r = p.add_run(item)
        set_run_font(r, size=BODY_SIZE)
    doc.add_page_break()


def build_report() -> None:
    ensure_dirs()
    arch = make_architecture_diagram()
    flow = make_message_flow_diagram()
    mobile_flow = make_mobile_workflow_diagram()
    db_design = make_database_diagram()
    logo = prepare_image("requirements434.png", max_width=500)

    imgs = {
        "login": prepare_image("requirements1649.png"),
        "login_filled": prepare_image("requirements1661.png"),
        "register": prepare_image("requirements1669.png"),
        "success": prepare_image("requirements1671.png"),
        "create_room": prepare_image("requirements1654.png"),
        "all_rooms": prepare_image("requirements1673.png"),
        "my_rooms": prepare_image("requirements1657.png"),
        "public_chat": prepare_image("requirements1667.png"),
        "users": prepare_image("requirements1663.png"),
        "contacts": prepare_image("requirements1660.png"),
        "db": prepare_image("requirements1682.png"),
        "server": prepare_image("requirements1690.png"),
        "client": prepare_image("requirements1692.png"),
    }

    doc = Document()
    configure_document(doc)
    add_cover(doc, logo)
    add_score_table(doc)
    add_toc(doc)

    add_heading(doc, "摘要", 1)
    add_para(doc, "本报告围绕“张善茹 Chat App”移动端实时聊天系统展开说明。该系统采用 H5 移动端浏览器方案实现，前端使用 React、TypeScript、Vite 和 Tailwind CSS 构建响应式界面，后端使用 Node.js、Express、Socket.io 和 Sequelize 提供接口、实时通信与数据持久化能力，数据库选用 SQLite。系统实现了用户注册登录、JWT 身份认证、聊天室创建与加入、公共聊天室实时消息、联系人管理、私聊、聊天记录持久化等功能，符合课程对前后端开发、数据库应用和文档说明的要求。")
    add_para(doc, "在实现过程中，我把重点放在两个方面：一是让移动端页面真的适合手机纵向浏览和单手操作，二是让消息在“实时显示”和“刷新后可恢复”之间保持一致。因此系统没有只停留在页面展示，而是将 Socket.io 事件与数据库记录结合起来，形成从用户操作到后端落库再到界面更新的完整闭环。")
    add_para(doc, "关键词：移动端应用；React；Socket.io；SQLite；实时聊天；数据持久化", first_line=False, bold=True)

    add_heading(doc, "1 项目概述", 1)
    add_heading(doc, "1.1 选题背景与目标", 2)
    add_para(doc, "即时通信是移动端应用中非常典型的场景，既涉及界面交互，也涉及账号认证、网络通信、消息同步和数据存储。相比只做静态页面，实时聊天系统更能综合体现移动终端程序设计课程中“前端界面、后端服务、数据库持久化、代码组织”几个核心要求。")
    add_para(doc, "本项目的目标是实现一个可以在手机浏览器中使用的实时聊天应用。用户登录后可以创建聊天室、加入已有聊天室、发送公共消息，也可以添加联系人并发起私聊。系统还需要保存用户、房间、联系人关系和聊天记录，使用户再次进入房间时能够看到历史消息。")
    add_heading(doc, "1.2 开发环境与技术栈", 2)
    env_rows = [
        ("前端框架", "React 19、TypeScript、Vite", "组件化页面、路由保护、状态管理和移动端界面渲染"),
        ("样式方案", "Tailwind CSS", "快速完成卡片、按钮、列表和聊天窗口的响应式样式"),
        ("实时通信", "Socket.io / socket.io-client", "完成加入房间、发送消息、接收消息等实时事件"),
        ("后端服务", "Node.js、Express、TypeScript", "提供 REST API、Socket.io 服务和业务控制器"),
        ("数据库", "SQLite、Sequelize", "保存账号、聊天室、联系人、公共消息和私聊消息"),
        ("安全认证", "JWT、bcrypt", "密码加密存储，登录后使用 token 访问受保护接口"),
    ]
    add_simple_table(doc, "表1.1 项目技术栈与作用", ("类别", "技术", "作用"), env_rows, widths=(3.0, 5.0, 8.0))
    add_heading(doc, "1.3 系统总体架构", 2)
    add_para(doc, "系统采用前后端分离结构。前端负责页面交互和状态展示，后端同时承载 REST API 与 Socket.io 实时连接，数据库通过 Sequelize 统一访问。REST API 更适合处理登录、注册、列表查询和历史消息读取；Socket.io 更适合处理聊天中的实时事件。这样拆分后，每个部分职责更清楚，也便于后续扩展文件消息、在线状态、消息已读等功能。")
    add_picture(doc, arch, 5.9, "图1.1 系统总体架构图")

    add_heading(doc, "2 需求分析与总体设计", 1)
    add_heading(doc, "2.1 功能需求", 2)
    req_rows = [
        ("用户认证", "注册、登录、JWT 保存、路由保护", "已实现"),
        ("聊天室", "创建房间、查看全部房间、加入房间、查看我的房间", "已实现"),
        ("实时聊天", "进入房间后发送和接收消息，消息按时间展示", "已实现"),
        ("联系人", "查找用户、添加联系人、查看联系人列表", "已实现"),
        ("私聊", "基于两名用户 ID 生成私聊房间并保存消息", "已实现"),
        ("持久化", "用户、房间、联系人、公共消息、私聊消息写入 SQLite", "已实现"),
    ]
    add_simple_table(doc, "表2.1 系统功能需求与实现情况", ("模块", "功能说明", "完成情况"), req_rows, widths=(3.0, 9.0, 3.0))
    add_heading(doc, "2.2 界面与交互设计", 2)
    add_para(doc, "界面设计采用浅色背景、白色卡片和琥珀色主按钮，尽量减少移动端的小字和复杂层级。登录/注册使用标签页切换，避免两个页面来回跳转；首页把 Rooms、Contacts、Find people 放在顶部导航，用户可以围绕“找房间、找联系人、开始聊天”三个动作操作。")
    add_para(doc, "我在界面上特别注意了按钮的触控面积和移动端纵向滚动体验。聊天室卡片中保留房间名和描述，用户不必进入房间就能判断是否加入；聊天区固定在页面主要区域，输入框和发送按钮放在底部，符合手机聊天软件的常见使用习惯。")
    add_picture(doc, mobile_flow, 5.9, "图2.1 移动端主要操作流程")
    add_heading(doc, "2.3 数据库设计", 2)
    add_para(doc, "数据库采用 SQLite，模型由 Sequelize 定义。主要实体包括用户、聊天室、联系人、公共消息和私聊消息。用户与聊天室之间是多对多关系，通过 RoomUser 中间表保存；联系人表中保存双向联系人关系；消息表分别保存公共聊天室消息和私聊消息。这样的设计使公共聊天和私聊逻辑既能复用统一的聊天窗口，又能在数据库中保持清晰边界。")
    add_picture(doc, db_design, 5.9, "图2.2 数据库关系设计图")
    add_simple_table(doc, "表2.2 SQLite 数据表设计说明", ("数据表", "用途", "主要字段与当前记录"), db_rows(), widths=(3.5, 5.0, 8.5))
    add_picture(doc, imgs["db"], 5.9, "图2.3 SQLite 中 users 表结构检查")

    add_heading(doc, "3 系统功能实现", 1)
    add_heading(doc, "3.1 用户注册与登录", 2)
    add_para(doc, "用户首次进入系统时会看到登录/注册页。注册时提交用户名、邮箱和密码，后端使用 bcrypt 对密码进行哈希处理后再写入数据库；登录时后端根据邮箱查询用户并校验密码，校验通过后生成 JWT 返回给前端。前端把 token 保存到 localStorage 中，并通过 ProtectedRoute 控制未登录用户不能访问 /home 页面。")
    add_picture_grid(doc, [(imgs["login"], "登录页"), (imgs["register"], "注册信息填写"), (imgs["success"], "注册成功反馈")], "图3.1 用户登录与注册界面", width_in=1.85)
    add_heading(doc, "3.2 聊天室创建、加入与切换", 2)
    add_para(doc, "进入首页后，用户可以创建新房间，也可以查看尚未加入的房间列表。创建房间时前端会把当前用户 ID 作为管理员 ID 一起提交，后端创建 Room 记录后把创建者加入 RoomUser 关系表。加入房间后，用户可以在 My Rooms 中看到自己的房间，并点击 Start Conversation 切换当前会话。")
    add_picture_grid(doc, [(imgs["create_room"], "创建房间"), (imgs["all_rooms"], "全部房间"), (imgs["my_rooms"], "我的房间")], "图3.2 聊天室创建、加入与我的房间", width_in=1.85)
    add_heading(doc, "3.3 公共聊天与私聊", 2)
    add_para(doc, "公共聊天使用房间 ID 作为 Socket.io room。用户点击某个房间后，前端调用 join_room 加入对应房间；发送消息时触发 send_message 事件，后端判断这是公共消息后写入 room_messages 表，并广播 receive_message。前端收到事件后将消息追加到 messageList 中，同时自动滚动到底部。")
    add_para(doc, "私聊部分没有单独开发另一个聊天窗口，而是复用了 ChatBox。系统将两名用户的 ID 排序后用下划线拼接成私聊 roomId，这样无论 A 找 B 还是 B 找 A，得到的私聊房间 ID 都相同，避免同一组用户出现两条会话记录。这是我在实现中觉得比较关键的一个小设计。")
    add_picture_grid(doc, [(imgs["public_chat"], "公共聊天室消息"), (imgs["users"], "用户列表与添加联系人"), (imgs["contacts"], "联系人私聊")], "图3.3 公共聊天、用户列表与私聊界面", width_in=1.85)
    add_heading(doc, "3.4 前后端运行效果", 2)
    add_para(doc, "本地运行时，后端监听 3030 端口，启动后会连接 SQLite 并同步检查 users、rooms、contacts、private_messages、room_messages、RoomUser 等表；前端通过 Vite 在 5173 端口启动。前端环境变量配置 API 地址和 Socket 地址，使开发时前后端可以独立运行。")
    add_picture_grid(doc, [(imgs["server"], "后端服务启动"), (imgs["client"], "前端 Vite 启动")], "图3.4 前后端启动与联调效果", width_in=2.85)

    add_heading(doc, "4 核心实现与代码说明", 1)
    add_heading(doc, "4.1 认证与路由保护", 2)
    add_para(doc, "身份认证由 REST 接口和前端上下文共同完成。后端注册接口负责密码哈希，登录接口负责密码比对并生成 token；前端 AuthProvider 读取 localStorage 中的 token 判断登录状态，并把 login/logout 方法提供给页面使用。这样处理后，用户刷新页面时仍可以保持登录状态，而退出登录时也能清除本地身份信息。")
    add_code_listing(doc, "代码清单4.1 登录成功后保存 token 并进入首页", """
if (response.token) {
    socket?.emit('register_user', response.userId);
    login(response.token);
    navigate('/home');
} else {
    setError('No token received');
}
""")
    add_heading(doc, "4.2 Socket.io 实时通信", 2)
    add_para(doc, "Socket 通信的关键在于房间机制。前端进入房间时发送 join_room，后端调用 socket.join(roomId)；发送消息时前端统一提交 roomId、content、senderId、isPrivate 等数据，后端根据 isPrivate 决定写入公共消息表还是私聊消息表，最后只向对应房间广播消息。")
    add_picture(doc, flow, 5.9, "图4.1 消息发送与持久化流程")
    add_algorithm(doc)
    add_code_listing(doc, "代码清单4.2 后端接收消息、写入数据库并广播", """
socket.on('send_message', async (messageData) => {
    const { roomId, content, senderId, isPrivate, receiverId, username } = messageData;

    if (isPrivate) {
        await PrivateMessage.create({ senderId, receiverId: receiverId!, content });
    } else {
        await RoomMessage.create({ senderId, roomId, content });
    }

    io.to(roomId).emit('receive_message', {
        content,
        username,
        senderId,
        isPrivate,
    });
});
""")
    add_heading(doc, "4.3 数据库模型与历史记录读取", 2)
    add_para(doc, "历史记录读取通过 REST 接口完成。公共聊天室根据 roomId 查询 room_messages，并关联 User 获取发送者用户名；私聊根据两个用户 ID 查询 senderId/receiverId 两个方向的消息，再按 createdAt 升序排列。前端进入某个房间时先请求历史记录，再监听实时消息，这样能兼顾“已有记录可追溯”和“新消息即时出现”。")
    add_code_listing(doc, "代码清单4.3 进入聊天室时加载历史消息", """
const endpoint = isPrivate
  ? `${url}/private/${parsedUser.id}/${roomId.split('_').find(id => id !== parsedUser.id)}`
  : `${url}/room_chat/${roomId}`;

const response = await fetch(endpoint, {
  method: 'GET',
  headers: {
    'Content-Type': 'application/json',
    'Authorization': `Bearer ${localStorage.getItem('token')}`
  }
});
const data = await response.json();
""")

    add_heading(doc, "5 测试结果与问题分析", 1)
    add_heading(doc, "5.1 测试环境", 2)
    add_para(doc, "测试环境为 Windows 本地开发环境，后端使用 npm start 启动，前端使用 npm run dev 启动，浏览器访问 http://localhost:5173。测试账号包括 juan.perez@example.com / Ju@nP3r3z! 和 maria.garcia@example.com / M@r1a2024，同时也通过注册页面新增账号进行验证。")
    test_rows = [
        ("注册账号", "输入用户名、邮箱、密码并提交", "页面提示注册成功，数据库 users 表新增记录", "通过"),
        ("登录认证", "使用测试账号登录", "返回 token，进入 /home，刷新后仍保持登录", "通过"),
        ("创建房间", "填写房间名和描述", "房间列表更新，创建者自动加入我的房间", "通过"),
        ("加入房间", "在全部房间中点击 Join", "房间从可加入列表转入我的房间", "通过"),
        ("公共聊天", "进入房间发送消息", "当前房间用户收到消息，刷新后历史消息仍存在", "通过"),
        ("添加联系人", "在用户列表中点击 Add", "联系人列表显示新增用户", "通过"),
        ("私聊", "从联系人列表发起会话并发送消息", "私聊消息写入 private_messages 表并实时显示", "通过"),
    ]
    add_simple_table(doc, "表5.1 功能测试结果", ("测试项", "操作", "预期与实际结果", "结论"), test_rows, widths=(3.0, 5.0, 7.0, 2.0))
    add_heading(doc, "5.2 调试问题与解决思路", 2)
    add_para(doc, "在调试中，最容易出错的是“状态更新”和“实时连接”的时序。例如 token 保存到 localStorage 后，Socket 连接也需要携带认证信息；如果用户刚登录就立刻连接，可能出现 token 还没有更新到 hook 状态中的情况。因此我在登录成功后显式保存 token，并在后续请求中统一从 localStorage 读取 Authorization。")
    add_para(doc, "另一个问题是私聊房间 ID 的一致性。如果直接用“当前用户_联系人”的顺序拼接，双方进入私聊时会得到两个不同 roomId，消息就无法汇合。最后采用排序后的 ID 拼接，让同一对用户永远落到同一个私聊房间中，这个处理虽然简单，但对聊天系统的数据一致性很重要。")
    add_para(doc, "界面方面，移动端高度有限，房间列表、联系人列表和聊天窗口都需要滚动。最初聊天消息多时会停留在旧位置，体验不自然；后来在 messageList 变化后主动把聊天容器滚动到底部，模拟常见聊天软件的阅读习惯。")

    add_heading(doc, "6 总结与心得", 1)
    add_heading(doc, "6.1 课程收获", 2)
    add_para(doc, "通过这个项目，我对移动端应用的理解从“页面能显示”推进到了“数据能流动、状态能保持、操作能闭环”。一个聊天按钮背后并不是简单的点击事件，而是表单状态、用户身份、Socket 连接、数据库写入、广播更新和历史记录读取之间的协作。这个过程让我更直观地理解了前后端分离项目中每一层的责任。")
    add_para(doc, "我也体会到，移动端界面设计不能只追求元素丰富，更重要的是让用户在小屏幕上少思考、少误触。比如登录/注册做成标签页，房间操作集中在卡片右侧，聊天输入区固定在底部，这些细节虽然不是复杂算法，但会直接影响作品是否“好用”。")
    add_heading(doc, "6.2 不足与改进", 2)
    add_para(doc, "目前系统已经完成课程要求的主要功能，但仍有改进空间。第一，消息内容现在主要是纯文本，后续可以扩展图片、文件和表情；第二，历史消息一次性读取，数据量变大后应增加分页加载；第三，在线状态还可以做得更准确，例如断线重连后自动同步；第四，可以增加消息撤回、未读数、搜索聊天记录等更贴近日常聊天软件的功能。")
    add_para(doc, "总体来看，这次结课项目让我把 React 组件化、Express 接口、Socket.io 事件和 SQLite 数据库真正串联起来。相比只完成某一个功能点，我更有收获的是理解了一个移动端应用从界面到数据的完整链路，也发现自己在调试异步状态和实时通信时还需要更多耐心和系统化思路。后续如果继续完善，我希望把它做成一个更接近真实产品的轻量级校园交流工具。")

    doc.save(OUT_FILE)
    print(str(OUT_FILE))


if __name__ == "__main__":
    build_report()
